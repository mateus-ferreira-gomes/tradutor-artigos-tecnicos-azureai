"""
📂 src/utils.py

Este módulo contém funções auxiliares do projeto, como manipulação de documentos Word
e integração com a função de tradução do Azure Translator.
"""

from docx import Document
from src.azure_translator import translator_text

# Idioma de destino padrão
language_destination = "pt-br"

def translate_document(path):
    """
    Traduz um documento Word inteiro para o idioma definido em language_destination
    usando a função translator_text do módulo azure_translator.
    
    Args:
        path (str): Caminho para o arquivo .docx de entrada.
    
    Returns:
        str: Caminho do arquivo traduzido.
    """
    # Abre o documento Word no caminho especificado
    document = Document(path)
    # Cria uma lista para armazenar os textos traduzidos
    full_text = []

    # Percorre todos os parágrafos do documento original
    for paragraph in document.paragraphs:
        # Traduz o texto de cada parágrafo
        translated_text = translator_text(paragraph.text, language_destination)
        # Adiciona o texto traduzido à lista
        full_text.append(translated_text)

    # Cria um novo documento Word vazio para armazenar a tradução
    translated_doc = Document()
    for line in full_text:
        print(line)  # imprime no console para conferência
        translated_doc.add_paragraph(line)

    # Define o nome do arquivo traduzido
    path_translated = path.replace(".docx", f"_{language_destination}.docx")
    # Salva o novo documento traduzido
    translated_doc.save(path_translated)

    return path_translated